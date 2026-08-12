from io import BytesIO
from openpyxl import Workbook, load_workbook
from docx import Document

FIELDS=["Setup Name","Description","Technologies","Keywords","Sources","Date Range","Start Date","End Date"]
def preview(data: bytes, ext: str):
    values={}
    if ext==".xlsx":
        ws=load_workbook(BytesIO(data),read_only=True,data_only=True).active
        for row in ws.iter_rows(values_only=True):
            if row and row[0] in FIELDS: values[row[0]]=row[1] if len(row)>1 else ""
    elif ext==".docx":
        doc=Document(BytesIO(data))
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells)>=2 and row.cells[0].text in FIELDS: values[row.cells[0].text]=row.cells[1].text
    else: raise ValueError("Only XLSX and DOCX files are supported.")
    if not values: raise ValueError("No recognized setup fields were found.")
    for key in ("Technologies","Keywords","Sources"):
        values[key]=[x.strip() for x in str(values.get(key,"")).split(",") if x.strip()]
    return values
def sample_xlsx():
    wb=Workbook(); ws=wb.active; ws.title="Setup Import"
    for i,(k,v) in enumerate({"Setup Name":"Sample Setup","Description":"Example monitoring setup","Technologies":"Windows 11, Outlook Web Access, Exchange Server","Keywords":"CVE, Exploit, RCE","Sources":"The Hacker News, CISA, Microsoft MSRC","Date Range":"7d"}.items(),1):
        ws.cell(i,1,k); ws.cell(i,2,v)
    out=BytesIO(); wb.save(out); return out.getvalue()
def sample_docx():
    doc=Document(); doc.add_heading("My ThreatLens Setup Import",0); table=doc.add_table(rows=0,cols=2)
    for k,v in {"Setup Name":"Sample Setup","Description":"Example monitoring setup","Technologies":"Windows 11, Outlook Web Access","Keywords":"CVE, Exploit","Sources":"The Hacker News, CISA","Date Range":"7d"}.items():
        cells=table.add_row().cells; cells[0].text=k; cells[1].text=v
    out=BytesIO(); doc.save(out); return out.getvalue()
